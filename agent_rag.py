import autogen
import ollama
from typing import List, Dict
import json
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader, PyPDFLoader
import os
import glob
import pandas as pd
from docx import Document
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import uuid
import time
from opencc import OpenCC

# 導入自定義的 JSONVectorDB
from vector_db import JSONVectorDB

def convert_to_traditional(text: str) -> str:
    """將簡體中文轉換為繁體中文"""
    cc = OpenCC('s2t')  # 簡體到繁體
    return cc.convert(text)

class CustomRAGAgentSystem:
    def __init__(self, reset_db=False, db_path="./custom_json_rag_db"):
        # Ollama配置
        self.base_url = "http://localhost:11434/v1"
        self.embedding_model = "tsd_4500datas_summary20250606_epoch11_f32:latest"
        self.llm_model = "qwen3:30b"
        # self.llm_model = "gemma3:27b"
        
        # 初始化自定義資料庫
        self.db_path = db_path
        if reset_db and os.path.exists(db_path):
            import shutil
            shutil.rmtree(db_path)
        
        # 創建自定義資料庫實例
        self.collection = JSONVectorDB(db_path)
        
        # 如果需要重置
        if reset_db:
            try:
                self.collection.delete_collection("rag_docs")
                print("已清除舊的RAG資料庫")
            except:
                pass
        
        # AutoGen配置
        self.config_list = [
            {
                'base_url': self.base_url,
                'api_key': "fakekey",
                'model': self.llm_model,
            }
        ]
        
        self.llm_config = {
            "config_list": self.config_list,
            "temperature": 0.0,
        }
        
        # 添加任務中斷標誌
        self._stop_flag = False
        
        self.setup_agents()
    
    def stop_current_task(self):
        """設置停止標誌"""
        self._stop_flag = True
        print("任務停止標誌已設置")
    
    def _check_stop_flag(self):
        """檢查是否應該停止任務"""
        if self._stop_flag:
            raise Exception("任務已被用戶中斷")
    
    def setup_agents(self):
        """設置AutoGen代理"""
        
        # 文檔篩選代理
        self.document_filter = autogen.AssistantAgent(
            name="DocumentFilter",
            llm_config=self.llm_config,
#             system_message="""您是文檔篩選專家。您的任務是：
# 1. 仔細閱讀提供的文檔內容
# 2. 判斷文檔是否與用戶問題相關
# 3. 如果相關，返回 "RELEVANT: [相關原因]"
# 4. 如果不相關，返回 "NOT_RELEVANT: [不相關原因]"
# 5. 請保持客觀和準確的判斷
# 請只返回判斷結果，不要添加其他內容。"""
#             system_message="""您是東擎科技(ASRock Industrial)技術支援部門(TSD)的專業文檔篩選專家。您的任務是:
# 1.仔細閱讀提供的文檔內容
# 2.精確判斷歷史技術支援文檔是否與用戶查詢相關。
#     **篩選標準（必須同時滿足）：**
#     **產品型號**：文檔中的產品型號必須與用戶詢問的型號一致，但有時候會有系列的問題可以以常理判斷
#     **型號識別規則**：
#     - 4x4-7XXX = 7000系列
#     - 4x4-6XXX = 6000系列  
#     - 4x4-5XXX = 5000系列
#     - 依此類推
#     **客戶**：文檔中的客戶必須與用戶查詢的問題匹配
#     **問題類型相關**：文檔中描述的技術問題必須與用戶查詢的問題類型相關
# 3. 如果相關，返回 "RELEVANT: [相關原因]"
# 4. 如果不相關，返回 "NOT_RELEVANT: [不相關原因]"
# 5.請只返回判斷結果，不要添加其他內容。"""
        system_message="""您是東擎科技(ASRock Industrial)技術支援部門(TSD)的專業文檔篩選專家。您的任務是:
1.仔細閱讀提供的文檔內容
2.精確判斷歷史技術支援文檔是否與用戶查詢相關。
3.客戶必須與用戶問題的客戶匹配
4.如果相關，返回 "RELEVANT: [相關原因]"
5.如果不相關，返回 "NOT_RELEVANT: [不相關原因]"
6.請只返回判斷結果，不要添加其他內容。

**型號識別規則**：
    - 4x4-7XXX = 7000系列
    - 4x4-6XXX = 6000系列  
    - 4x4-5XXX = 5000系列
    - 依此類推
    - NUC MTL = NUC 125 155
"""
        )
        
        # 答案整合代理
        self.answer_synthesizer = autogen.AssistantAgent(
            name="AnswerSynthesizer",
            llm_config=self.llm_config,
            system_message="""您是答案整合專家。您的任務是：
1. 基於篩選後的相關文檔，為用戶問題提供綜合性答案
2. 整合所有相關信息，提供完整且準確的回答
3. 如果信息不足，請明確指出
4. 使用繁體中文回答
請提供詳細、有用的答案。"""
        )
    
    def add_documents_from_directory(self, directory_path: str, file_patterns: List[str] = None):
        """從資料夾載入所有文檔到RAG資料庫"""
        if file_patterns is None:
            file_patterns = [
                "*.txt", "*.pdf", "*.md", "*.json",
                "*.docx", "*.doc",  # Word文檔
                "*.xlsx", "*.xls"   # Excel文檔
            ]
        
        print(f"開始載入資料夾: {directory_path}")
        print(f"支援的文件格式: {', '.join(file_patterns)}")
        
        if not os.path.exists(directory_path):
            print(f"錯誤: 資料夾 {directory_path} 不存在")
            return False
        
        total_files = 0
        successful_files = 0
        
        for pattern in file_patterns:
            file_path_pattern = os.path.join(directory_path, pattern)
            files = glob.glob(file_path_pattern)
            
            for file_path in files:
                total_files += 1
                file_ext = os.path.splitext(file_path)[1].lower()
                
                try:
                    print(f"處理文件: {os.path.basename(file_path)}")
                    
                    if file_ext == '.pdf':
                        success = self.add_document(file_path, doc_type="pdf")
                    elif file_ext in ['.txt', '.md']:
                        success = self.add_document(file_path, doc_type="txt")
                    elif file_ext == '.json':
                        success = self.add_json_document(file_path)
                    elif file_ext in ['.docx', '.doc']:
                        success = self.add_word_document(file_path)
                    elif file_ext in ['.xlsx', '.xls']:
                        success = self.add_excel_document(file_path)
                    else:
                        print(f"跳過不支援的文件類型: {file_path}")
                        continue
                    
                    if success:
                        successful_files += 1
                        
                except Exception as e:
                    print(f"處理文件 {file_path} 時發生錯誤: {e}")
        
        print(f"\n載入完成: 成功處理 {successful_files}/{total_files} 個文件")
        return successful_files > 0
    
    def add_word_document(self, file_path: str, original_filename: str = None):
        """處理Word文檔 (.docx, .doc)"""
        try:
            print(f"正在處理Word文檔: {file_path}")
            
            doc = Document(file_path)
            
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    full_text.append("表格內容:\n" + "\n".join(table_text))
            
            text = "\n\n".join(full_text)
            
            if not text.strip():
                print(f"警告: Word文檔 {file_path} 沒有文本內容")
                return False
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=850,
                chunk_overlap=100,
                length_function=len,
            )
            chunks = text_splitter.split_text(text)
            
            ids = []
            embeddings = []
            documents = []
            metadatas = []
            
            for i, chunk in enumerate(chunks):
                response = ollama.embed(model=self.embedding_model, input=chunk)
                embedding = response["embeddings"][0]
                
                doc_id = f"{os.path.basename(file_path)}_{i}"
                ids.append(doc_id)
                embeddings.append(embedding)
                documents.append(chunk)
                metadatas.append({
                    "source": file_path, 
                    "chunk_id": i, 
                    "file_type": "docx",
                    "original_filename": original_filename if original_filename else os.path.basename(file_path),
                    "paragraphs_count": len(doc.paragraphs),
                    "tables_count": len(doc.tables)
                })
            
            self.collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas
            )
            
            print(f"成功添加Word文檔: {file_path}，共{len(chunks)}個片段")
            print(f"  - 段落數: {len(doc.paragraphs)}")
            print(f"  - 表格數: {len(doc.tables)}")
            return True
            
        except Exception as e:
            print(f"添加Word文檔失敗: {e}")
            return False
    
    def add_excel_document(self, file_path: str, original_filename: str = None):
        """處理Excel文檔 (.xlsx, .xls)"""
        try:
            print(f"正在處理Excel文檔: {file_path}")
            
            try:
                excel_file = pd.ExcelFile(file_path)
                sheet_names = excel_file.sheet_names
                print(f"  - 發現 {len(sheet_names)} 個工作表: {sheet_names}")
                
                all_text = []
                
                for sheet_name in sheet_names:
                    print(f"  - 處理工作表: {sheet_name}")
                    
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    if df.empty:
                        print(f"    工作表 {sheet_name} 為空，跳過")
                        continue
                    
                    sheet_text = [f"工作表: {sheet_name}"]
                    sheet_text.append(f"行數: {len(df)}, 列數: {len(df.columns)}")
                    sheet_text.append("列名: " + " | ".join(str(col) for col in df.columns))
                    sheet_text.append("")
                    
                    max_rows = min(1000, len(df))
                    for i in range(max_rows):
                        row_data = []
                        for col in df.columns:
                            cell_value = df.iloc[i][col]
                            if pd.isna(cell_value):
                                cell_value = ""
                            else:
                                cell_value = str(cell_value).strip()
                            row_data.append(cell_value)
                        
                        if any(row_data):
                            sheet_text.append(" | ".join(row_data))
                    
                    if len(df) > max_rows:
                        sheet_text.append(f"... (還有 {len(df) - max_rows} 行數據)")
                    
                    all_text.append("\n".join(sheet_text))
                
                text = "\n\n" + "="*50 + "\n\n".join(all_text)
                
            except Exception as e:
                print(f"讀取Excel文件時發生錯誤: {e}")
                return False
            
            if not text.strip():
                print(f"警告: Excel文檔 {file_path} 沒有數據內容")
                return False
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1200,
                chunk_overlap=100,
                length_function=len,
            )
            chunks = text_splitter.split_text(text)
            
            ids = []
            embeddings = []
            documents = []
            metadatas = []
            
            for i, chunk in enumerate(chunks):
                response = ollama.embed(model=self.embedding_model, input=chunk)
                embedding = response["embeddings"][0]
                
                doc_id = f"{os.path.basename(file_path)}_{i}"
                ids.append(doc_id)
                embeddings.append(embedding)
                documents.append(chunk)
                metadatas.append({
                    "source": file_path, 
                    "chunk_id": i, 
                    "file_type": "excel",
                    "original_filename": original_filename if original_filename else os.path.basename(file_path),
                    "sheets_count": len(sheet_names),
                    "sheet_names": sheet_names
                })
            
            self.collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas
            )
            
            print(f"成功添加Excel文檔: {file_path}，共{len(chunks)}個片段")
            print(f"  - 工作表數: {len(sheet_names)}")
            return True
            
        except Exception as e:
            print(f"添加Excel文檔失敗: {e}")
            return False
    
    def add_json_document(self, file_path: str, original_filename: str = None):
        """專門處理JSON文檔"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
            
            if isinstance(json_data, dict):
                text = json.dumps(json_data, ensure_ascii=False, indent=2)
            elif isinstance(json_data, list):
                text = "\n".join([json.dumps(item, ensure_ascii=False) for item in json_data])
            else:
                text = str(json_data)
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=850,
                chunk_overlap=0,
                length_function=len,
            )
            chunks = text_splitter.split_text(text)
            
            ids = []
            embeddings = []
            documents = []
            metadatas = []
            
            for i, chunk in enumerate(chunks):
                response = ollama.embed(model=self.embedding_model, input=chunk)
                embedding = response["embeddings"][0]
                
                doc_id = f"{os.path.basename(file_path)}_{i}"
                ids.append(doc_id)
                embeddings.append(embedding)
                documents.append(chunk)
                metadatas.append({
                    "source": file_path, 
                    "chunk_id": i, 
                    "file_type": "json",
                    "original_filename": original_filename if original_filename else os.path.basename(file_path)
                })
            
            self.collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas
            )
            
            print(f"成功添加JSON文檔: {file_path}，共{len(chunks)}個片段")
            return True
            
        except Exception as e:
            print(f"添加JSON文檔失敗: {e}")
            return False
    
    def add_document(self, file_path: str, doc_type: str = "txt", original_filename: str = None):
        """添加單個文檔到RAG資料庫"""
        try:
            if doc_type.lower() == "pdf":
                loader = PyPDFLoader(file_path)
                documents = loader.load()
                text = "\n".join([doc.page_content for doc in documents])
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            
            # 檢查是否包含時間戳模板
            has_timestamp_template = "編號與日期:" in text
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=850,
                chunk_overlap=0,
                length_function=len,
            )
            chunks = text_splitter.split_text(text)
            
            ids = []
            embeddings = []
            documents = []
            metadatas = []
            
            for i, chunk in enumerate(chunks):
                response = ollama.embed(model=self.embedding_model, input=chunk)
                embedding = response["embeddings"][0]
                
                doc_id = f"{os.path.basename(file_path)}_{i}"
                ids.append(doc_id)
                embeddings.append(embedding)
                documents.append(chunk)
                metadatas.append({
                    "source": file_path, 
                    "chunk_id": i, 
                    "file_type": doc_type,
                    "original_filename": original_filename if original_filename else os.path.basename(file_path),
                    "has_timestamp_template": has_timestamp_template
                })
            
            self.collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas
            )
            
            print(f"成功添加文檔: {file_path}，共{len(chunks)}個片段")
            if has_timestamp_template:
                print(f"  - 包含時間戳模板")
            return True
            
        except Exception as e:
            print(f"添加文檔失敗: {e}")
            return False
    
    def list_loaded_documents(self):
        """顯示已載入的文檔列表"""
        try:
            all_docs = self.collection.get(include=["metadatas"])
            
            file_stats = {}
            for metadata in all_docs["metadatas"]:
                source = metadata.get("source", "unknown")
                filename = os.path.basename(source)
                file_type = metadata.get("file_type", "unknown")
                
                if filename not in file_stats:
                    file_stats[filename] = {
                        "chunks": 0,
                        "type": file_type,
                        "path": source,
                        "extra_info": {}
                    }
                file_stats[filename]["chunks"] += 1
                
                if file_type == "docx":
                    file_stats[filename]["extra_info"]["paragraphs"] = metadata.get("paragraphs_count", 0)
                    file_stats[filename]["extra_info"]["tables"] = metadata.get("tables_count", 0)
                elif file_type == "excel":
                    file_stats[filename]["extra_info"]["sheets"] = metadata.get("sheets_count", 0)
                    file_stats[filename]["extra_info"]["sheet_names"] = metadata.get("sheet_names", [])
            
            print(f"\n已載入的文檔列表 (共 {len(file_stats)} 個文件):")
            for filename, stats in file_stats.items():
                extra_info = ""
                if stats["type"] == "docx" and stats["extra_info"]:
                    extra_info = f" [段落:{stats['extra_info'].get('paragraphs', 0)}, 表格:{stats['extra_info'].get('tables', 0)}]";
                elif stats["type"] == "excel" and stats["extra_info"]:
                    sheets = stats['extra_info'].get('sheets', 0)
                    extra_info = f" [工作表:{sheets}]"
                
                icon = {
                    "txt": "📄", "md": "📝", "pdf": "📕", 
                    "json": "📋", "docx": "📘", "excel": "📊"
                }.get(stats["type"], "📄")
                
                print(f"  {icon} {stats['display_name']} ({stats['type']}) - {stats['chunks']} 個片段{extra_info}")
                
        except Exception as e:
            print(f"獲取文檔列表失敗: {e}")
    
    def search_documents(self, query: str, n_results: int = 6, date_range: str = '') -> List[Dict]:
        """搜索最相關的文檔"""
        try:
            print(f"開始搜索，查詢: {query}, 時間區間: {date_range}")
            
            # 檢查停止標誌
            self._check_stop_flag()
            
            # 獲取查詢的嵌入向量
            response = ollama.embed(model=self.embedding_model, input=query)
            query_embedding = np.array(response["embeddings"][0]).reshape(1, -1)
            print(f"查詢向量維度: {query_embedding.shape}")
            
            # 檢查停止標誌
            self._check_stop_flag()
            
            # 獲取所有文檔
            all_docs = self.collection.get(include=["documents", "metadatas", "embeddings"])
            print(f"獲取到 {len(all_docs['documents'])} 個文檔")
            
            # 調試：檢查數據結構
            print("文檔結構:", {
                "documents": len(all_docs["documents"]),
                "metadatas": len(all_docs["metadatas"]),
                "embeddings": len(all_docs["embeddings"])
            })
            
            # 如果有時間區間，進行過濾
            if date_range and date_range.strip():
                try:
                    # 檢查是否為 "all time"
                    if date_range.strip() == "all time":
                        print("選擇了所有時間範圍，不進行時間過濾")
                        filtered_docs = []
                        for i in range(len(all_docs["documents"])):
                            doc_content = all_docs["documents"][i]
                            doc_metadata = all_docs["metadatas"][i]
                            doc_embedding = all_docs["embeddings"][i]
                            doc_timestamp = None
                            
                            # 嘗試從元數據中獲取時間戳
                            if "timestamp" in doc_metadata:
                                try:
                                    doc_timestamp = int(doc_metadata["timestamp"])
                                except (ValueError, TypeError):
                                    pass
                            
                            # 如果元數據中沒有，則檢查文檔內容
                            if doc_timestamp is None and "編號與日期:" in doc_content:
                                import re
                                pattern = r'\(\d+\)\s*(\d{8})'
                                match = re.search(pattern, doc_content)
                                if match:
                                    doc_timestamp = int(match.group(1))
                            
                            filtered_docs.append({
                                "document": doc_content,
                                "metadata": doc_metadata,
                                "embedding": doc_embedding,
                                "timestamp": str(doc_timestamp) if doc_timestamp else ""
                            })
                    else:
                        start_date, end_date = date_range.split(' - ')
                        start_date = int(start_date.strip())
                        end_date = int(end_date.strip())
                        print(f"時間區間: {start_date} - {end_date}")
                        
                        filtered_docs = []
                        for i in range(len(all_docs["documents"])):
                            # 從文檔內容中提取時間戳
                            doc_content = all_docs["documents"][i]
                            doc_metadata = all_docs["metadatas"][i]
                            doc_embedding = all_docs["embeddings"][i]
                            doc_timestamp = None
                            
                            # 首先檢查元數據中是否有時間戳
                            if "timestamp" in doc_metadata:
                                try:
                                    doc_timestamp = int(doc_metadata["timestamp"])
                                except (ValueError, TypeError) as e:
                                    print(f"解析時間戳失敗: {e}, 原始值: {doc_metadata['timestamp']}")
                                    pass
                            
                            # 如果元數據中沒有，則檢查文檔內容
                            if doc_timestamp is None and "編號與日期:" in doc_content:
                                import re
                                pattern = r'\(\d+\)\s*(\d{8})'
                                match = re.search(pattern, doc_content)
                                if match:
                                    doc_timestamp = int(match.group(1))
                                    print(f"從內容中找到時間戳: {doc_timestamp}")
                            
                            # 如果還是沒有找到時間戳，跳過此文檔
                            if doc_timestamp is None:
                                print(f"文檔沒有時間戳，跳過: {doc_metadata.get('original_filename', 'unknown')}")
                                continue
                            
                            # 檢查時間戳是否在指定範圍內
                            if start_date <= doc_timestamp <= end_date:
                                filtered_docs.append({
                                    "document": doc_content,
                                    "metadata": doc_metadata,
                                    "embedding": doc_embedding,
                                    "timestamp": str(doc_timestamp)
                                })
                        
                        print(f"時間區間過濾後剩餘 {len(filtered_docs)} 個文檔")
                        
                        if len(filtered_docs) == 0:
                            print("沒有找到符合時間區間的文檔")
                            return []
                            
                except Exception as e:
                    print(f"時間區間過濾失敗: {str(e)}")
                    # 如果時間過濾失敗，則回退到不過濾，但仍嘗試提取時間戳
                    filtered_docs = []
                    for i in range(len(all_docs["documents"])):
                        doc_content = all_docs["documents"][i]
                        doc_metadata = all_docs["metadatas"][i]
                        doc_embedding = all_docs["embeddings"][i]
                        doc_timestamp = None
                        if "timestamp" in doc_metadata:
                            try:
                                doc_timestamp = int(doc_metadata["timestamp"])
                            except (ValueError, TypeError):
                                pass
                        if doc_timestamp is None and "編號與日期:" in doc_content:
                            import re
                            pattern = r'\(\d+\)\s*(\d{8})'
                            match = re.search(pattern, doc_content)
                            if match:
                                doc_timestamp = int(match.group(1))
                        
                        filtered_docs.append({
                            "document": doc_content,
                            "metadata": doc_metadata,
                            "embedding": doc_embedding,
                            "timestamp": str(doc_timestamp) if doc_timestamp else ""
                        })
            else:
                # 如果沒有時間區間，直接使用所有文檔，但仍然提取時間戳
                filtered_docs = []
                for i in range(len(all_docs["documents"])):
                    doc_content = all_docs["documents"][i]
                    doc_metadata = all_docs["metadatas"][i]
                    doc_embedding = all_docs["embeddings"][i]
                    doc_timestamp = None
                    
                    # 嘗試從元數據中獲取時間戳
                    if "timestamp" in doc_metadata:
                        try:
                            doc_timestamp = int(doc_metadata["timestamp"])
                        except (ValueError, TypeError):
                            pass
                    
                    # 如果元數據中沒有，則檢查文檔內容
                    if doc_timestamp is None and "編號與日期:" in doc_content:
                        import re
                        pattern = r'\(\d+\)\s*(\d{8})'
                        match = re.search(pattern, doc_content)
                        if match:
                            doc_timestamp = int(match.group(1))
                    
                    filtered_docs.append({
                        "document": doc_content,
                        "metadata": doc_metadata,
                        "embedding": doc_embedding,
                        "timestamp": str(doc_timestamp) if doc_timestamp else "" # 確保有timestamp字段
                    })
            
            # 計算相似度
            results = []
            for doc in filtered_docs:
                try:
                    # 確保向量維度正確
                    doc_vec = np.array(doc["embedding"]).reshape(1, -1)
                    if doc_vec.shape[1] != query_embedding.shape[1]:
                        print(f"向量維度不匹配: 文檔向量 {doc_vec.shape}, 查詢向量 {query_embedding.shape}")
                        continue
                        
                    similarity = cosine_similarity(query_embedding, doc_vec)[0][0]
                    distance = 1 - similarity
                    results.append({
                        "content": doc["document"],
                        "metadata": doc["metadata"],
                        "distance": distance,
                        "chunk_id": doc["metadata"].get("chunk_id", 0),
                        "timestamp": doc.get("timestamp", "")
                    })
                except Exception as e:
                    print(f"計算相似度時發生錯誤: {str(e)}")
                    continue
            
            # 按距離排序並返回前 n_results 個結果
            results.sort(key=lambda x: x["distance"])
            final_results = results[:n_results]
            print(f"返回 {len(final_results)} 個結果")
            
            # 在關鍵步驟後檢查停止標誌
            self._check_stop_flag()
            
            return final_results
            
        except Exception as e:
            if self._stop_flag:
                print("搜索任務被用戶中斷")
                return []
            print(f"文檔搜索失敗: {str(e)}")
            return []
    
    def filter_documents(self, query: str, documents: List[Dict]) -> List[Dict]:
        """使用第一個LLM篩選文檔"""
        relevant_docs = []
        filter_interactions = {}  # 保存每個文檔的交互訊息
        
        for doc in documents:
            # 檢查停止標誌
            self._check_stop_flag()
            
            filter_user_proxy = autogen.UserProxyAgent(
                name="filter_user_proxy",
                human_input_mode="NEVER",
                max_consecutive_auto_reply=1,
                is_termination_msg=lambda x: True,
                code_execution_config=False,
                llm_config=self.llm_config,
            )
            
            filter_prompt = f"""
用戶問題: {query}

文檔內容:
{doc['content']}

請判斷此文檔是否與用戶問題相關。
"""
            
            # 記錄交互開始
            interaction_messages = []
            interaction_messages.append({
                'role': 'user',
                'content': filter_prompt,
                'timestamp': time.time()
            })
            
            filter_user_proxy.initiate_chat(
                self.document_filter,
                message=filter_prompt,
                max_turns=1
            )
            
            try:
                last_message = filter_user_proxy.last_message(self.document_filter)
                if last_message:
                    # 將AI分析結果轉換為繁體中文
                    traditional_content = convert_to_traditional(last_message["content"])
                    
                    interaction_messages.append({
                        'role': 'assistant',
                        'content': traditional_content,
                        'timestamp': time.time()
                    })
                    
                    # 保存交互訊息
                    doc_key = f"{doc['metadata'].get('original_filename', os.path.basename(doc['metadata']['source']))}_{doc['metadata'].get('chunk_id', 0)}"
                    filter_interactions[doc_key] = {
                        'messages': interaction_messages,
                        'is_relevant': "NOT_RELEVANT:" not in traditional_content,
                        'filename': doc['metadata'].get('original_filename', os.path.basename(doc['metadata']['source'])),
                        'chunk_id': doc['metadata'].get('chunk_id', 0)
                    }
                    
                    if "NOT_RELEVANT:" not in traditional_content:
                        relevant_docs.append(doc)
            except Exception as e:
                print(f"獲取篩選結果失敗: {e}")
                # 即使失敗也保存交互訊息
                doc_key = f"{doc['metadata'].get('original_filename', os.path.basename(doc['metadata']['source']))}_{doc['metadata'].get('chunk_id', 0)}"
                filter_interactions[doc_key] = {
                    'messages': interaction_messages,
                    'is_relevant': True,  # 失敗時默認認為相關
                    'filename': doc['metadata'].get('original_filename', os.path.basename(doc['metadata']['source'])),
                    'chunk_id': doc['metadata'].get('chunk_id', 0),
                    'error': str(e)
                }
                relevant_docs.append(doc)
            
            # 每處理完一個文檔後休息3秒
            time.sleep(12)
            
            # 檢查停止標誌
            self._check_stop_flag()
        
        # 將交互訊息添加到每個相關文檔中
        for doc in relevant_docs:
            doc_key = f"{doc['metadata'].get('original_filename', os.path.basename(doc['metadata']['source']))}_{doc['metadata'].get('chunk_id', 0)}"
            if doc_key in filter_interactions:
                doc['filter_interaction'] = filter_interactions[doc_key]
        
        return relevant_docs, filter_interactions
    
    def generate_answer(self, query: str, relevant_docs: List[Dict]) -> str:
        """使用第二個LLM生成最終答案"""
        if len(relevant_docs) == 0:
            return "沒有找到相關文檔來回答您的問題。"
        
        # 檢查停止標誌
        self._check_stop_flag()
        
        synthesis_user_proxy = autogen.UserProxyAgent(
            name="synthesis_user_proxy",
            human_input_mode="NEVER",
            max_consecutive_auto_reply=1,
            is_termination_msg=lambda x: x.get("content", "").rstrip().endswith("TERMINATE"),
            code_execution_config=False,
            llm_config=self.llm_config,
        )
        
        context = "\n\n".join([f"文檔{i+1}:\n{doc['content']}" 
                              for i, doc in enumerate(relevant_docs)])
        
        synthesis_prompt = f"""
基於以下相關文檔回答用戶問題:

用戶問題: {query}

相關文檔:
{context}

請提供綜合性的答案:
"""
        
        synthesis_user_proxy.initiate_chat(
            self.answer_synthesizer,
            message=synthesis_prompt,
            max_turns=1
        )
        
        try:
            # 檢查停止標誌
            self._check_stop_flag()
            
            last_message = synthesis_user_proxy.last_message(self.answer_synthesizer)
            if last_message:
                return convert_to_traditional(last_message["content"])
                # return last_message["content"]
            return "生成答案失敗"
        except Exception as e:
            if self._stop_flag:
                return "答案生成被用戶中斷"
            print(f"獲取答案失敗: {e}")
            return "生成答案過程中發生錯誤"
    
    def query(self, question: str) -> str:
        """完整的RAG查詢流程"""
        print(f"\n開始處理問題: {question}")
        
        print("1. 搜索相關文檔...")
        documents = self.search_documents(question, n_results=6)
        print(f"找到 {len(documents)} 個候選文檔")
        
        if not documents:
            return "沒有找到相關文檔"
        
        print("找到的文檔:")
        for i, doc in enumerate(documents, 1):
            filename = os.path.basename(doc['metadata']['source'])
            file_type = doc['metadata'].get('file_type', 'unknown')
            distance = doc['distance']
            print(f"  {i}. {filename} ({file_type}) - 相似度: {1-distance:.3f}")
        
        print("\n2. 篩選相關文檔...")
        relevant_docs, filter_interactions = self.filter_documents(question, documents)
        print(f"篩選後保留 {len(relevant_docs)} 個相關文檔")
        
        if relevant_docs:
            print("保留的文檔:")
            for i, doc in enumerate(relevant_docs, 1):
                filename = os.path.basename(doc['metadata']['source'])
                file_type = doc['metadata'].get('file_type', 'unknown')
                print(f"  {i}. {filename} ({file_type})")
        
        print("\n3. 生成最終答案...")
        answer = self.generate_answer(question, relevant_docs)
        
        return answer 